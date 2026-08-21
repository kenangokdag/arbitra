"""Plan: docs/plans/RAPOR_DOCX_EXPORT_2026-08-16.md §4.1.

Coverage: build_docx() v1 ve v2 raporlardan geçerli .docx bytes üretiyor mu —
python-docx ile geri okuyup beklenen içeriğin (başlık/verdict/boyut/finding/
risk-radar/atıf sayaçları) dokümanda VAR olduğunu doğrular.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone
from uuid import uuid4

import pytest
from docx import Document as DocxDocument

from api.models.review import (
    ActionItem,
    CitationIntegritySummary,
    DimensionScore,
    EvidencePack,
    ExecutiveVerdict,
    Finding,
    ManuscriptAnchor,
    ManuscriptMeta,
    ReviewerCouncilItem,
    ReviewProvenance,
    ReviewReport,
    RiskRadarItem,
    SectionReview,
)
from api.services.report_export_service import build_docx

pytestmark = pytest.mark.unit


def _provenance() -> ReviewProvenance:
    return ReviewProvenance(
        model_used="gemini-flash-tr",
        persona_version="v1",
        engine_version="v1",
        generated_at=datetime.now(timezone.utc),
    )


def _all_text(docx_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_build_docx_v1_report_contains_core_fields() -> None:
    report = ReviewReport(
        mode="author",
        language="tr",
        manuscript_meta=ManuscriptMeta(title="Örnek Makale Başlığı"),
        summary="Bu makale özet metnidir.",
        overall_assessment="Genel değerlendirme metni.",
        verdict="major_revision",
        final_score=6.5,
        dimension_scores=[
            DimensionScore(key="soundness", score=6.0, rationale="Yöntem orta düzeyde."),
        ],
        evidence_pack=EvidencePack(
            citation_integrity=CitationIntegritySummary(
                total=5, resolved=4, fabricated=1, retracted=0, not_found_in_index=0
            ),
        ),
        provenance=_provenance(),
    )

    docx_bytes = build_docx(report, uuid4())
    assert docx_bytes.startswith(b"PK")  # docx = zip paketi, dürüst sanity check
    text = _all_text(docx_bytes)

    assert "Örnek Makale Başlığı" in text
    assert "Büyük revizyon" in text  # _humanize_verdict("major_revision")
    assert "Yöntem sağlamlığı" in text  # _humanize_dimension("soundness")
    assert "Bu makale özet metnidir." in text
    assert "Genel değerlendirme metni." in text
    assert "5 kaynak" in text
    assert "çelişkili: 1" in text


def test_build_docx_v2_report_contains_findings_and_risk_radar() -> None:
    report = ReviewReport(
        mode="author",
        language="tr",
        manuscript_meta=ManuscriptMeta(title="V2 Makale"),
        summary="özet",
        overall_assessment="genel",
        verdict="reject",
        provenance=_provenance(),
        executive_verdict=ExecutiveVerdict(
            overall_readiness_score=30.0,
            recommended_decision="reject",
            one_sentence_diagnosis="Ciddi yöntem sorunları var.",
            top_fatal_risks=["Örneklem çok küçük"],
        ),
        risk_radar=[
            RiskRadarItem(
                dimension="soundness",
                score=20.0,
                severity="critical",
                why_it_matters="Sonuçlar güvenilir değil.",
            ),
        ],
        findings=[
            Finding(
                finding_id="F-1",
                dimension="soundness",
                severity="critical",
                title="Örneklem büyüklüğü yetersiz",
                summary="n=5 ile genelleme yapılamaz.",
                global_issue=True,
                action_item_ids=["A-1"],
                manuscript_anchors=[
                    ManuscriptAnchor(anchor_id="methods.p2", section="Yöntem", quote="n=5 katılımcı ile..."),
                ],
            ),
        ],
        action_plan=[
            ActionItem(
                action_id="A-1",
                priority="P0",
                instruction="Örneklem büyüklüğünü artırın veya sınırlılık olarak tartışın.",
            ),
        ],
        reviewer_council=[
            ReviewerCouncilItem(
                role="methodologist",
                summary="Yöntem ciddi şekilde zayıf.",
                key_objection="Örneklem gerekçesi yok.",
            ),
        ],
        section_reviews=[
            SectionReview(section="Yöntem", status="broken", what_breaks="Örneklem gerekçesi eksik."),
        ],
    )

    docx_bytes = build_docx(report, uuid4())
    text = _all_text(docx_bytes)

    assert "Ciddi yöntem sorunları var." in text
    assert "Örneklem çok küçük" in text
    assert "Örneklem büyüklüğü yetersiz" in text
    assert "n=5 ile genelleme yapılamaz." in text
    assert "n=5 katılımcı ile" in text  # manuscript anchor quote
    assert "Örneklem büyüklüğünü artırın" in text  # linked action item
    assert "Yöntem ciddi şekilde zayıf." in text  # reviewer council
    assert "Örneklem gerekçesi eksik." in text  # section review
