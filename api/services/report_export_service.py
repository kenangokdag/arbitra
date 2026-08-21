"""Hakem raporu → .docx export (RAPOR_DOCX_EXPORT_2026-08-16).

Plan: docs/plans/RAPOR_DOCX_EXPORT_2026-08-16.md §2.1
Girdi: ReviewReport (api/models/review.py) — aynı model GET /{job_id}/report'un
döndürdüğü model. Çıktı: python-docx ile üretilmiş .docx bytes.

python-docx ZATEN proje bağımlılığı (pyproject.toml:41, ingestion S1'de .docx
OKUMAK için) — burada YAZMAK için reuse ediliyor, yeni paket yok.

Dürüst sınır: web/src/components/review/ReviewReportView.tsx'in interaktif
3-katmanlı sunumunun (collapse/drawer) birebir kopyası DEĞİL — sade, okunabilir,
profesyonel bir Word dokümanı. Şiddet (severity) düz metin etiketi olarak yazılır
("[KRİTİK]"), renkli rozet değil.

isV2 mantığı ReviewReportView.tsx:84-93 ile PARALEL ama AYRI yazıldı (TS/Python
iki runtime, paylaşılan kod imkânsız) — plan §2.1'de bilinçli kod-tekrarı olarak
not edildi.
"""

from __future__ import annotations

from io import BytesIO
from uuid import UUID

from docx import Document
from docx.shared import Pt

from api.models.review import (
    ActionItem,
    Finding,
    ManuscriptAnchor,
    ReviewerCouncilItem,
    ReviewReport,
    RiskRadarItem,
)

VERDICT_LABELS_TR: dict[str, str] = {
    "accept": "Kabul",
    "minor_revision": "Küçük revizyon",
    "major_revision": "Büyük revizyon",
    "reject": "Ret",
}

DIMENSION_LABELS_TR: dict[str, str] = {
    "originality": "Özgünlük",
    "importance": "Önem",
    "claims_supported": "İddia desteği",
    "soundness": "Yöntem sağlamlığı",
    "clarity": "Anlaşılırlık",
    "community_value": "Alana katkı",
    "contextualization": "Bağlamlandırma",
    "citation_integrity": "Atıf bütünlüğü",
    "coverage_completeness": "Kapsam tamlığı",
    "statistical_consistency": "İstatistik tutarlılığı",
}


def _humanize_dimension(key: str) -> str:
    return DIMENSION_LABELS_TR.get(key, key)


def _humanize_verdict(key: str) -> str:
    return VERDICT_LABELS_TR.get(key, key)


def _is_v2(report: ReviewReport) -> bool:
    """ReviewReportView.tsx:84-93 ile PARALEL — v2 alanlarından biri doluysa v2."""
    return (
        report.schema_version.startswith("review_report.v2")
        or report.executive_verdict is not None
        or report.disclosure is not None
        or report.document_classification is not None
        or len(report.risk_radar) > 0
        or len(report.reviewer_council) > 0
        or len(report.findings) > 0
        or len(report.action_plan) > 0
        or len(report.section_reviews) > 0
    )


def build_docx(report: ReviewReport, job_id: UUID) -> bytes:
    """ReviewReport → .docx bytes. Kapsam: docs/plans/RAPOR_DOCX_EXPORT_2026-08-16.md."""
    doc = Document()

    title = report.manuscript_meta.title or "Başlıksız makale"
    doc.add_heading(title, level=0)

    decision = (
        report.executive_verdict.recommended_decision
        if report.executive_verdict is not None
        else report.verdict
    )
    p = doc.add_paragraph()
    p.add_run("Hakem kararı: ").bold = True
    p.add_run(_humanize_verdict(decision))
    if report.final_score is not None:
        p.add_run(f"  ·  Skor: {report.final_score:.1f} / 10")

    if _is_v2(report):
        _add_v2_sections(doc, report)
    else:
        _add_v1_sections(doc, report)

    _add_citation_integrity_section(doc, report)
    _add_provenance_section(doc, report, job_id)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_v1_sections(doc: Document, report: ReviewReport) -> None:
    doc.add_heading("Özet", level=1)
    doc.add_paragraph(report.summary)

    if report.dimension_scores:
        doc.add_heading("Boyut değerlendirmesi", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Boyut", "Skor", "Gerekçe"
        for d in report.dimension_scores:
            row = table.add_row().cells
            row[0].text = _humanize_dimension(d.key)
            row[1].text = f"{d.score:.1f} / 10"
            row[2].text = d.rationale

    if report.strengths:
        doc.add_heading("Güçlü yönler", level=1)
        for group in report.strengths:
            doc.add_heading(group.category, level=2)
            for point in group.points:
                doc.add_paragraph(point, style="List Bullet")

    if report.weaknesses:
        doc.add_heading("Geliştirilecek yönler", level=1)
        for group in report.weaknesses:
            doc.add_heading(group.category, level=2)
            for point in group.points:
                doc.add_paragraph(point, style="List Bullet")

    if report.detailed_comments:
        doc.add_heading("Detaylı yorumlar", level=1)
        for c in report.detailed_comments:
            p = doc.add_paragraph()
            p.add_run(f"{c.area}: ").bold = True
            p.add_run(c.comment)

    if report.questions:
        doc.add_heading("Yazara sorular", level=1)
        for q in report.questions:
            doc.add_paragraph(q, style="List Number")

    doc.add_heading("Genel değerlendirme", level=1)
    doc.add_paragraph(report.overall_assessment)


def _add_v2_sections(doc: Document, report: ReviewReport) -> None:
    ev = report.executive_verdict
    if ev is not None:
        doc.add_heading("Yönetici özeti", level=1)
        doc.add_paragraph(ev.one_sentence_diagnosis)
        p = doc.add_paragraph()
        p.add_run("Hazırlık puanı: ").bold = True
        p.add_run(f"{ev.overall_readiness_score:.0f} / 100")
        if ev.top_fatal_risks:
            doc.add_heading("En kritik riskler", level=2)
            for risk in ev.top_fatal_risks:
                doc.add_paragraph(risk, style="List Bullet")

    doc.add_heading("Özet", level=1)
    doc.add_paragraph(report.summary)

    if report.risk_radar:
        doc.add_heading("Risk radarı", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
            "Boyut",
            "Skor",
            "Önem",
            "Neden önemli",
        )
        for item in _sorted_by_severity(report.risk_radar):
            row = table.add_row().cells
            row[0].text = _humanize_dimension(item.dimension)
            row[1].text = "Değerlendirilmedi" if item.score is None else f"{item.score:.0f} / 100"
            row[2].text = item.severity.upper()
            row[3].text = item.why_it_matters

    if report.findings:
        doc.add_heading("Bulgular", level=1)
        action_by_id = {a.action_id: a for a in report.action_plan}
        for f in _sorted_by_severity(report.findings):
            _add_finding(doc, f, action_by_id)

    if report.reviewer_council:
        doc.add_heading("Hakem heyeti", level=1)
        for item in report.reviewer_council:
            _add_council_item(doc, item)

    if report.section_reviews:
        doc.add_heading("Bölüm bazında değerlendirme", level=1)
        for s in report.section_reviews:
            p = doc.add_paragraph()
            p.add_run(f"{s.section} [{s.status.upper()}]").bold = True
            if s.what_works:
                doc.add_paragraph(f"İyi: {s.what_works}")
            if s.what_breaks:
                doc.add_paragraph(f"Sorun: {s.what_breaks}")

    doc.add_heading("Genel değerlendirme", level=1)
    doc.add_paragraph(report.overall_assessment)


def _sorted_by_severity(items: list) -> list:
    """severity alanı olan herhangi bir liste — critical→info sıralı."""
    rank = {"critical": 0, "major": 1, "moderate": 2, "minor": 3, "info": 4}
    return sorted(items, key=lambda x: rank.get(x.severity, 99))


def _add_finding(doc: Document, finding: Finding, action_by_id: dict[str, ActionItem]) -> None:
    p = doc.add_paragraph()
    p.add_run(f"[{finding.severity.upper()}] {finding.title}").bold = True
    if finding.summary:
        doc.add_paragraph(finding.summary)
    for anchor in finding.manuscript_anchors:
        _add_anchor(doc, anchor)
    for action_id in finding.action_item_ids:
        action = action_by_id.get(action_id)
        if action is not None:
            fix_p = doc.add_paragraph()
            fix_p.add_run("Önerilen düzeltme: ").italic = True
            fix_p.add_run(action.instruction)


def _add_anchor(doc: Document, anchor: ManuscriptAnchor) -> None:
    if not anchor.quote:
        return
    quote_p = doc.add_paragraph()
    quote_p.paragraph_format.left_indent = Pt(18)
    run = quote_p.add_run(f"“{anchor.quote}”")
    run.italic = True


def _add_council_item(doc: Document, item: ReviewerCouncilItem) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{item.role}").bold = True
    if item.summary:
        doc.add_paragraph(item.summary)
    if item.key_objection:
        obj_p = doc.add_paragraph()
        obj_p.add_run("Ana itiraz: ").italic = True
        obj_p.add_run(item.key_objection)


def _add_citation_integrity_section(doc: Document, report: ReviewReport) -> None:
    ci = report.evidence_pack.citation_integrity
    doc.add_heading("Atıf bütünlüğü", level=1)
    doc.add_paragraph(
        f"Toplam {ci.total} kaynak — doğrulandı: {ci.resolved}, "
        f"indekste bulunamadı: {ci.not_found_in_index}, "
        f"çelişkili: {ci.fabricated}, geri çekilmiş: {ci.retracted}."
    )


def _add_provenance_section(doc: Document, report: ReviewReport, job_id: UUID) -> None:
    p = report.provenance
    doc.add_heading("Künye", level=2)
    doc.add_paragraph(f"İş kimliği: {job_id}")
    doc.add_paragraph(f"Model: {p.model_used}")
    doc.add_paragraph(f"Motor sürümü: {p.engine_version}")
    doc.add_paragraph(f"Üretildi: {p.generated_at.isoformat()}")
